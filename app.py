from flask import Flask, render_template, request, send_file
import os, uuid
import pandas as pd
from docx import Document
from docx2pdf import convert
from io import BytesIO
from zipfile import ZipFile

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/merge', methods=['POST'])
def merge():
    docx_file = request.files['docx']
    xls_file = request.files['xls']
    output_type = request.form['output_type']

    docx_path = os.path.join(UPLOAD_FOLDER, str(uuid.uuid4()) + ".docx")
    xls_path = os.path.join(UPLOAD_FOLDER, str(uuid.uuid4()) + ".xls")
    docx_file.save(docx_path)
    xls_file.save(xls_path)

    df = pd.read_excel(xls_path)
    files = []

    def fill_template(row):
        doc = Document(docx_path)
        for p in doc.paragraphs:
            for col in df.columns:
                if f'{{{{{col}}}}}' in p.text:
                    p.text = p.text.replace(f'{{{{{col}}}}}', str(row[col]))
        return doc

    if output_type == 'separate_docx':
        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, 'w') as zipf:
            for i, row in df.iterrows():
                doc = fill_template(row)
                temp_path = os.path.join(OUTPUT_FOLDER, f'merge_{i}.docx')
                doc.save(temp_path)
                zipf.write(temp_path, arcname=f'merge_{i}.docx')
        zip_buffer.seek(0)
        return send_file(zip_buffer, as_attachment=True, download_name='merged_docx.zip')

    elif output_type == 'one_docx':
        merged = Document()
        for _, row in df.iterrows():
            doc = fill_template(row)
            for element in doc.element.body:
                merged.element.body.append(element)
        path = os.path.join(OUTPUT_FOLDER, 'merged.docx')
        merged.save(path)
        return send_file(path, as_attachment=True)

    elif output_type == 'separate_pdf':
        temp_dir = os.path.join(OUTPUT_FOLDER, str(uuid.uuid4()))
        os.makedirs(temp_dir, exist_ok=True)
        for i, row in df.iterrows():
            doc = fill_template(row)
            temp_docx = os.path.join(temp_dir, f'temp_{i}.docx')
            temp_pdf = os.path.join(temp_dir, f'merge_{i}.pdf')
            doc.save(temp_docx)
            convert(temp_docx, temp_pdf)
        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, 'w') as zipf:
            for f in os.listdir(temp_dir):
                if f.endswith('.pdf'):
                    zipf.write(os.path.join(temp_dir, f), arcname=f)
        zip_buffer.seek(0)
        return send_file(zip_buffer, as_attachment=True, download_name='merged_pdf.zip')

    elif output_type == 'one_pdf':
        from PyPDF2 import PdfMerger
        merger = PdfMerger()
        temp_dir = os.path.join(OUTPUT_FOLDER, str(uuid.uuid4()))
        os.makedirs(temp_dir, exist_ok=True)
        for i, row in df.iterrows():
            doc = fill_template(row)
            temp_docx = os.path.join(temp_dir, f'temp_{i}.docx')
            temp_pdf = os.path.join(temp_dir, f'temp_{i}.pdf')
            doc.save(temp_docx)
            convert(temp_docx, temp_pdf)
            merger.append(temp_pdf)
        final_pdf = os.path.join(temp_dir, 'merged.pdf')
        merger.write(final_pdf)
        merger.close()
        return send_file(final_pdf, as_attachment=True)

    return "Invalid option", 400
